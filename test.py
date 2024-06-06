from docx import Document

def analyze_document(docx_file):
    document = Document(docx_file)
    methods = analyze_methods(document)
    chosen_method = determine_hiding_method(methods)
    encoded_text = encode_text(document, chosen_method)
    decoded_results = decode_text(encoded_text)
    print_results(decoded_results, chosen_method)


def analyze_methods(document):
    methods = {'color': False, 'font_size': False, 'spacing': False, 'scale': False}
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.font.color.rgb != (0, 0, 0):
                methods['color'] = True
            if run.font.size != 152400:
                methods['font_size'] = True
            if run_get_scale(run):
                methods['scale'] = True
            if run_get_spacing(run):
                methods['spacing'] = True
                break  # Оптимизация: если нашли пробел, прекратить поиск
    return methods


def determine_hiding_method(methods):
    return next((method for method, value in methods.items() if value), None)


def run_get_spacing(run):
    return bool(run._r.xpath(".//w:spacing"))


def run_get_scale(run):
    return bool(run._r.xpath(".//w:w"))


def encode_text(document, chosen_method):
    encoded_text = ""
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if chosen_method == 'color':
                encoded_text += "1" if run.font.color.rgb != (1, 1, 1) else "0"
            elif chosen_method == 'font_size':
                encoded_text += "1" if run.font.size != 152400 else "0"
            elif chosen_method == 'spacing':
                encoded_text += "1" if run_get_spacing(run) else "0"
            elif chosen_method == 'scale':
                encoded_text += "1" if run_get_scale(run) else "0"
            else:
                encoded_text += "0"  # If method is not defined
    return encoded_text


def decode_text(encoded_text):
    decoded_results = []
    chunks = [chunk for chunk in encoded_text.split("000000000") if chunk and chunk != '000000000']
    for chunk in chunks:
        chunk += "0" * (-len(chunk) % 8)
        binary_sequence = " ".join([chunk[i:i+8] for i in range(0, len(chunk), 8)])[:-1]
        decoded_results.extend(decode_text_for_encodings(binary_sequence))
    return decoded_results


def print_results(decoded_results, chosen_method):
    for binary_sequence, encoding, decoded_text in decoded_results:
        if decoded_text:
            print(f"Encoding Method: {encoding}, Binary Sequence: {binary_sequence}, Decoded Text: {decoded_text}")
    print(f"Chosen hiding method: {chosen_method if chosen_method else 'not determined'}")


def decode_text_for_encodings(binary_sequence):
    decoded_results = []
    encodings = ["koi8-r", "koi8-u", "cp866", "windows-1251", "cp1251"]
    for encoding in encodings:
        binary_sequence, decoded_text = decode_bytes_to_string(binary_sequence, encoding)
        if decoded_text:
            decoded_results.append((binary_sequence, encoding, decoded_text))
    return decoded_results


def decode_bytes_to_string(binary_sequence, encoding):
    try:
        bytes_list = [int(byte, 2) for byte in binary_sequence.split()]
        bytes_data = bytes(bytes_list)
        decoded_text = bytes_data.decode(encoding)
        return binary_sequence, decoded_text
    except Exception as e:
        return None, None

# Example usage
analyze_document("variant17.docx")
